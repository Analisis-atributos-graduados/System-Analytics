import logging
import sys
import os
import io
import uuid
import json
import re
import traceback
import zipfile
from typing import List
import urllib.parse
import hmac
import hashlib
import fitz
import base64
import pdfplumber
from docx import Document
from datetime import datetime, timezone

import google.auth
import google.auth.transport.requests

try:
    from logging_config import setup_logging
    setup_logging()
except ImportError:
    logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

log = logging.getLogger(__name__)
log.info("--- INICIO DE ARRANQUE DEL WORKER ---")

log.info("Importando: Librerías...")
import requests
from PyPDF2 import PdfReader, PdfWriter
from thefuzz import process as fuzzy_process
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph
from reportlab.lib.units import inch
from docx import Document

log.info("Importando: FastAPI y componentes SQL...")
from fastapi import FastAPI, Request, HTTPException, Depends, File, UploadFile, Form, Body
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy.orm import Session, joinedload
import sqlalchemy.exc

log.info("Importando: Google Cloud Services...")
from google.cloud import storage, tasks_v2
from google.cloud import iam_credentials_v1

try:
    from google import genai
    from google.genai.types import Part
    log.info("✓ Gemini client importado")
except ImportError:
    log.warning("⚠ google-genai no disponible")
    genai = None

log.info("Importando: Módulos locales...")
from database import get_db, engine, Base
from models import Evaluacion, ArchivoProcesado, CriterioConfig, ResultadoAnalisis
from schemas import (
    CriterioConfigUpdate, ExamBatchRequest, GenerateUploadURLRequest,
    FileTaskPayload, EvaluationTaskPayload, EvaluacionDetailSchema, EvaluacionSchema,
)

log.info("--- ¡IMPORTACIONES COMPLETADAS! ---")
log.info("Configurando variables de entorno...")
GCP_PROJECT_ID = os.environ.get("GCP_PROJECT_ID")
GCP_LOCATION = os.environ.get("GCP_LOCATION")
BUCKET_NAME = os.environ.get("GCS_BUCKET_NAME")
QUEUE_NAME = os.environ.get("GCS_QUEUE_NAME")
RAPIDAPI_KEY = os.environ.get("RAPIDAPI_KEY")
SERVICE_URL = os.environ.get("SERVICE_URL")

if not SERVICE_URL:
    SERVICE_URL = "https://analitica-backend-511391059179.southamerica-east1.run.app"

log.info(f"GCP_PROJECT_ID: {GCP_PROJECT_ID}")
log.info(f"GCP_LOCATION: {GCP_LOCATION}")
log.info(f"BUCKET_NAME: {BUCKET_NAME}")
log.info(f"QUEUE_NAME: {QUEUE_NAME}")
log.info(f"SERVICE_URL: {SERVICE_URL}")

log.info("Inicializando clientes de Google Cloud...")
storage_client = storage.Client(project=GCP_PROJECT_ID)
tasks_client = tasks_v2.CloudTasksClient()
iam_client = iam_credentials_v1.IAMCredentialsClient()

try:
    credentials, project = google.auth.default()
    log.info(f"✓ Credenciales obtenidas para proyecto: {project}")
except Exception as e:
    log.error(f"Error obteniendo credenciales: {e}")
    credentials = None

if genai:
    try:
        gemini_client = genai.Client(
            vertexai=True,
            project=GCP_PROJECT_ID,
            location="us-central1"
        )
        log.info("✓ Cliente Gemini inicializado")
    except Exception as e:
        log.warning(f"No se pudo inicializar Gemini: {e}")
        gemini_client = None
else:
    gemini_client = None

log.info("Inicializando FastAPI...")
app = FastAPI(title="EvalIA Backend")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

Base.metadata.create_all(bind=engine)
log.info("✓ Tablas de base de datos verificadas/creadas")

# Variables globales para modelo ML
modelo_ml_lazy = None
tokenizer_lazy = None

def get_classifier():
    """Carga el modelo DeBERTa para inferencia directa (no zero-shot)."""
    global modelo_ml_lazy, tokenizer_lazy

    if modelo_ml_lazy is None:
        log.info("LAZY LOADING: Cargando DeBERTa-v3-base...")
        try:
            from transformers import AutoTokenizer, AutoModelForSequenceClassification
            import torch

            # Cargar tokenizer
            tokenizer_lazy = AutoTokenizer.from_pretrained("/app/model")

            # Cargar modelo
            modelo_ml_lazy = AutoModelForSequenceClassification.from_pretrained("/app/model")
            modelo_ml_lazy.eval()  # Modo evaluación (más rápido)

            log.info("✓ DeBERTa cargado exitosamente")

        except Exception as e:
            log.exception(f"Error cargando DeBERTa: {e}")
            modelo_ml_lazy = None
            tokenizer_lazy = None

    return modelo_ml_lazy, tokenizer_lazy

log.info("--- ARRANQUE COMPLETADO ---")


# ============================================================================
# FUNCIONES AUXILIARES: Extracción de imágenes (para ensayos/informes)
# ============================================================================

def extraer_texto_word(file_bytes: bytes) -> str:
    """Extrae texto de un archivo Word (.docx)."""
    try:
        doc = Document(io.BytesIO(file_bytes))
        texto = "\n\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
        log.info(f"Texto extraído de Word: {len(texto)} caracteres")
        return texto
    except Exception as e:
        log.exception(f"Error extrayendo texto de Word: {e}")
        return ""

def extract_images_from_pdf(file_bytes: bytes) -> list:
    """Extrae todas las imágenes de un PDF. Retorna lista de bytes."""
    images = []
    try:
        pdf = fitz.open(stream=file_bytes, filetype="pdf")
        for page_num in range(len(pdf)):
            page = pdf[page_num]
            image_list = page.get_images(full=True)
            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = pdf.extract_image(xref)
                image_bytes = base_image["image"]
                images.append(image_bytes)
                log.info(f"Extraída imagen {img_index + 1} de página {page_num + 1}")
        pdf.close()
        log.info(f"Total de imágenes extraídas del PDF: {len(images)}")
    except Exception as e:
        log.exception(f"Error extrayendo imágenes del PDF: {e}")
    return images


def extract_images_from_docx(file_bytes: bytes) -> list:
    """Extrae todas las imágenes de un DOCX. Retorna lista de bytes."""
    images = []
    try:
        doc = Document(io.BytesIO(file_bytes))
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_bytes = rel.target_part.blob
                images.append(image_bytes)
                log.info(f"Extraída imagen de DOCX: {rel.target_ref}")
        log.info(f"Total de imágenes extraídas del DOCX: {len(images)}")
    except Exception as e:
        log.exception(f"Error extrayendo imágenes del DOCX: {e}")
    return images


def analyze_images_with_gemini(images: list, tema: str, descripcion_tema: str) -> dict:
    """Analiza imágenes/diagramas usando Gemini Vision."""
    if not gemini_client:
        log.error("Gemini client no disponible")
        return {"error": "Gemini no disponible"}

    if not images:
        log.warning("No hay imágenes para analizar")
        return {"analisis_imagenes": "No se encontraron imágenes"}

    try:
        prompt = f"""Eres un evaluador experto en análisis de documentos académicos.

Analiza las imágenes/diagramas presentes en este documento sobre el tema: "{tema}"
Descripción: {descripcion_tema}

Evalúa los siguientes aspectos:

1. **Claridad Visual**: ¿Los diagramas/gráficos son claros y legibles?
2. **Relevancia al Tema**: ¿Las imágenes son pertinentes al tema?
3. **Calidad Técnica**: ¿Los diagramas tienen etiquetas, leyendas y están bien estructurados?
4. **Aporte al Contenido**: ¿Las imágenes complementan la comprensión del texto?

Devuelve tu análisis en formato JSON:
{{
  "claridad_visual": "descripción y puntaje 0-10",
  "relevancia_tema": "descripción y puntaje 0-10",
  "calidad_tecnica": "descripción y puntaje 0-10",
  "aporte_contenido": "descripción y puntaje 0-10",
  "comentarios_generales": "resumen general"
}}"""

        content_parts = [prompt]
        for idx, img_bytes in enumerate(images[:5]):
            mime_type = "image/jpeg"
            if img_bytes[:4] == b'\x89PNG':
                mime_type = "image/png"
            content_parts.append(Part.from_bytes(data=img_bytes, mime_type=mime_type))
            log.info(f"Añadida imagen {idx + 1} al análisis Gemini")

        response = gemini_client.models.generate_content(
            model="gemini-2.0-flash-exp",
            contents=content_parts
        )

        response_text = response.text.strip()
        log.info(f"Respuesta raw de Gemini: {response_text[:200]}...")

        # Attempt to extract JSON from a markdown code block
        json_match = re.search(r'```json\s*(.*?)\s*```', response_text, re.DOTALL)
        if json_match:
            response_text = json_match.group(1).strip()
        else:
            # If not a specific 'json' block, try a generic code block
            json_match = re.search(r'```\s*(.*?)\s*```', response_text, re.DOTALL)
            if json_match:
                response_text = json_match.group(1).strip()

        try:
            analisis = json.loads(response_text)
            log.info(f"Análisis Gemini exitoso")
            return analisis
        except json.JSONDecodeError as json_err:
            log.error(f"Error parseando JSON: {json_err}")
            return {
                "error": "No se pudo parsear JSON",
                "comentarios_generales": response_text[:500]
            }

    except Exception as e:
        log.exception(f"Error en análisis Gemini: {e}")
        return {"error": f"Error en análisis visual: {str(e)}"}


# ============================================================================
# FUNCIONES AUXILIARES: Detección de tipo y extracción de texto
# ============================================================================

def detectar_tipo_y_extraer_texto(file_bytes: bytes, filename: str) -> dict:
    """
    Detecta si el documento tiene texto extraíble.
    Para ensayos/informes: extrae texto + imágenes.
    Para exámenes manuscritos: solo detecta que requiere OCR.
    """
    resultado = {
        "tipo": "desconocido",
        "texto": "",
        "tiene_imagenes": False,
        "imagenes": [],
        "requiere_ocr": False
    }

    filename_lower = filename.lower()

    try:
        if filename_lower.endswith('.pdf'):
            import pdfplumber
            pdf = pdfplumber.open(io.BytesIO(file_bytes))
            texto_completo = ""

            for page in pdf.pages:
                texto_pagina = page.extract_text() or ""
                texto_completo += texto_pagina

            pdf.close()

            imagenes_pdf = extract_images_from_pdf(file_bytes)
            resultado["imagenes"] = imagenes_pdf
            resultado["tiene_imagenes"] = len(imagenes_pdf) > 0

            texto_limpio = texto_completo.strip()
            num_caracteres = len(texto_limpio)

            if num_caracteres > 200:
                resultado["tipo"] = "texto_extraible"
                resultado["texto"] = texto_limpio
                resultado["requiere_ocr"] = False
                log.info(f"PDF con texto extraíble: {num_caracteres} chars, {len(imagenes_pdf)} imágenes")
            else:
                resultado["tipo"] = "manuscrito"
                resultado["requiere_ocr"] = True
                log.info(f"PDF manuscrito detectado, requiere OCR")

        elif filename_lower.endswith('.docx'):
            doc = Document(io.BytesIO(file_bytes))
            texto_completo = "\n".join([p.text for p in doc.paragraphs])
            imagenes_docx = extract_images_from_docx(file_bytes)
            resultado["imagenes"] = imagenes_docx
            resultado["tiene_imagenes"] = len(imagenes_docx) > 0
            resultado["tipo"] = "texto_extraible"
            resultado["texto"] = texto_completo.strip()
            resultado["requiere_ocr"] = False
            log.info(f"DOCX detectado: {len(texto_completo)} chars, {len(imagenes_docx)} imágenes")

        else:
            log.warning(f"Tipo de archivo no soportado: {filename}")
            resultado["tipo"] = "no_soportado"

    except Exception as e:
        log.exception(f"Error detectando tipo: {e}")
        resultado["tipo"] = "error"
        resultado["requiere_ocr"] = True

    return resultado


# ============================================================================
# FUNCIÓN AUXILIAR: Búsqueda de nombres con fuzzy matching
# ============================================================================

def find_student_name_in_text(ocr_text: str, student_list: list) -> str:
    """Busca nombre del alumno en texto OCR con fuzzy matching."""
    if not ocr_text or not student_list:
        return None

    found_direct = []

    patterns = [
        r"(?:nombres? y apellidos|apellidos y nombres?|alumno|estudiante)\s*[:\-\s]\s*([a-zA-Z\sÁÉÍÓÚáéíóúñÑ,'\. ]+)",
        r"^(?:nombre|alumno|estudiante)[:\s]+([a-zA-Z\sÁÉÍÓÚáéíóúñÑ,'\. ]+)"
    ]

    text_lines = ocr_text.strip().split("\n")
    potential_names = []

    for line in text_lines:
        line_stripped = line.strip()
        if not line_stripped:
            continue
        for pattern in patterns:
            match = re.search(pattern, line_stripped, re.IGNORECASE)
            if match:
                name_found = match.group(1).strip()
                name_found = re.sub(r'[^\w\sÁÉÍÓÚáéíóúñÑ]+$', '', name_found).strip()
                if name_found:
                    log.info(f"Nombre potencial por patrón: '{name_found}'")
                    potential_names.append(name_found)

    for i, line in enumerate(text_lines):
        line_lower = line.strip().lower()
        if any(line_lower.startswith(keyword) or keyword + ":" in line_lower
               for keyword in ["alumno", "nombre", "estudiante"]):
            parts = line.split(":", 1)
            if len(parts) > 1 and parts[1].strip():
                name_found = parts[1].strip()
                log.info(f"Nombre en línea de keyword: '{name_found}'")
                potential_names.append(name_found)
            elif i + 1 < len(text_lines):
                next_line = text_lines[i + 1].strip()
                if next_line and len(next_line) > 3:
                    log.info(f"Nombre en línea siguiente: '{next_line}'")
                    potential_names.append(next_line)

    if not potential_names:
        log.info("No se encontraron nombres por patrones, buscando directamente...")
        ocr_text_lower = ' '.join(ocr_text.lower().split())
        for student in student_list:
            student_lower = student.lower()
            if re.search(r'\b' + re.escape(student_lower) + r'\b', ocr_text_lower):
                log.info(f"Coincidencia directa: '{student}'")
                found_direct.append(student)

        if found_direct:
            best = max(found_direct, key=len)
            log.info(f"Mejor coincidencia directa: '{best}'")
            return best
        else:
            log.warning("No se encontró nombre")
            return None

    log.info(f"Nombres potenciales: {potential_names}")

    cleaned = set(re.sub(r'[^\w\sÁÉÍÓÚáéíóúñÑ]', '', name).strip() for name in potential_names if len(name.strip()) > 2)
    if not cleaned:
        return None

    best_query = max(cleaned, key=len)
    log.info(f"Query fuzzy: '{best_query}'")

    best_match = fuzzy_process.extractOne(best_query, student_list)

    if best_match and best_match[1] > 75:
        log.info(f"Fuzzy match: '{best_match[0]}' score {best_match[1]}")
        return best_match[0]
    else:
        log.warning(f"Fuzzy insuficiente: {best_match}")
        return None


# ============================================================================
# FUNCIONES AUXILIARES: OCR y análisis con modelo ML
# ============================================================================

def call_ocr_api(file_bytes: bytes, mime_type: str = "application/pdf") -> str:
    """Llama a la API de OCR de Pen-to-Print."""
    if not RAPIDAPI_KEY:
        log.error("RAPIDAPI_KEY no configurada")
        return ""

    url = "https://pen-to-print-handwriting-ocr.p.rapidapi.com/recognize/"

    headers = {
        "X-RapidAPI-Key": RAPIDAPI_KEY,
        "X-RapidAPI-Host": "pen-to-print-handwriting-ocr.p.rapidapi.com"
    }

    files = {
        "srcImg": ("document", file_bytes, mime_type),
        "Session": (None, "string")
    }

    try:
        log.info("Llamando a API de OCR...")
        response = requests.post(url, headers=headers, files=files, timeout=120)
        response.raise_for_status()
        data = response.json()

        text = data.get("value", "")
        log.info(f"OCR exitoso, longitud: {len(text)}")
        return text
    except Exception as e:
        log.exception(f"Error en OCR API: {e}")
        return ""


def analizar_criterio_especifico(texto: str, criterio: str) -> dict:
    """
    Analiza un criterio específico del texto usando DeBERTa.

    Args:
        texto: Texto a analizar
        criterio: Puede ser 'aplicacion_conceptos', 'relacion_contextual', 'coherencia_logica'

    Returns:
        dict con nivel, confidence, predicted_class
    """
    try:
        import torch

        model, tokenizer = get_classifier()
        if model is None or tokenizer is None:
            log.error("Modelo no disponible")
            return crear_resultado_error()

        # ✅ PROMPTS ESPECÍFICOS POR CRITERIO
        prompts = {
            "aplicacion_conceptos": f"Evalúa la aplicación de conceptos técnicos y teóricos en este texto académico: {texto[:2500]}",
            "relacion_contextual": f"Evalúa la relación contextual, conexión de ideas y relevancia del contenido: {texto[:2500]}",
            "coherencia_logica": f"Evalúa la coherencia lógica, estructura argumentativa y fluidez del texto: {texto[:2500]}"
        }

        prompt = prompts.get(criterio, texto[:2500])

        # Tokenizar con límite de 512 tokens
        inputs = tokenizer(
            prompt,
            max_length=512,
            truncation=True,
            padding=True,
            return_tensors="pt"
        )

        # Inferencia sin calcular gradientes (más rápido)
        with torch.no_grad():
            outputs = model(**inputs)
            logits = outputs.logits
            probs = torch.softmax(logits, dim=-1)

            predicted_class = torch.argmax(probs, dim=-1).item()
            confidence = probs[0][predicted_class].item()

        # Mapear a niveles de calidad
        class_labels = {
            0: "Excelente",
            1: "Bueno",
            2: "Regular"
        }

        nivel = class_labels.get(predicted_class, "Regular")

        log.info(f"  Criterio '{criterio}': {nivel} (confianza: {confidence:.2f})")

        return {
            "nivel": nivel,
            "confidence": float(confidence),
            "predicted_class": predicted_class
        }

    except Exception as e:
        log.exception(f"Error analizando criterio {criterio}: {e}")
        return crear_resultado_error()


def analizar_texto_con_modelo(texto_completo: str, analisis_visual: dict = None) -> dict:
    """
    Analiza el texto completo evaluando cada criterio por separado.

    Args:
        texto_completo: Texto del documento
        analisis_visual: Resultado del análisis de Gemini (opcional)

    Returns:
        dict con puntajes por criterio
    """
    log.info(f"📊 Analizando texto de {len(texto_completo)} caracteres...")

    try:
        # ✅ ANALIZAR CADA CRITERIO POR SEPARADO
        criterios = ["aplicacion_conceptos", "relacion_contextual", "coherencia_logica"]
        resultados_criterios = {}

        for criterio in criterios:
            log.info(f"Analizando criterio: {criterio}...")
            resultado = analizar_criterio_especifico(texto_completo, criterio)
            resultados_criterios[criterio] = resultado

        # ✅ CONVERTIR CADA RESULTADO A PUNTAJE (0-20)
        nivel_to_score = {
            "Excelente": 17.0,
            "Bueno": 10.0,
            "Regular": 3.0
        }

        puntajes_criterios = {}

        for criterio, resultado in resultados_criterios.items():
            nivel = resultado.get("nivel", "Regular")
            confidence = resultado.get("confidence", 0.5)

            puntaje_base = nivel_to_score.get(nivel, 10.0)
            ajuste = (confidence - 0.5) * 6  # ±3 puntos
            puntaje_ajustado = max(0, min(20, puntaje_base + ajuste))

            # Convertir a proporción (0-1)
            proporcion = puntaje_ajustado / 20

            puntajes_criterios[criterio] = {
                "nivel": nivel,
                "confidence": confidence,
                "puntaje": puntaje_ajustado,
                "proporcion": proporcion  # 0-1 para BD
            }

            log.info(f"  {criterio}: {nivel} → {puntaje_ajustado:.2f}/20 ({proporcion:.3f})")

        # ✅ INCORPORAR ANÁLISIS VISUAL SI EXISTE
        if analisis_visual and isinstance(analisis_visual, dict):
            # Ajustar puntajes basándose en calidad de imágenes
            claridad = analisis_visual.get("claridad_visual", {})
            relevancia = analisis_visual.get("relevancia_tema", {})

            if isinstance(claridad, dict):
                calidad_visual = claridad.get("puntaje", 5)
            else:
                calidad_visual = 5

            if isinstance(relevancia, dict):
                relevancia_puntaje = relevancia.get("puntaje", 5)
            else:
                relevancia_puntaje = 5

            # Factor de ajuste visual (0.8 a 1.2)
            factor_visual = (calidad_visual + relevancia_puntaje) / 10  # Promedio 0-10 → 0-1
            factor_visual = 0.8 + (factor_visual * 0.4)  # Rango: 0.8 a 1.2

            log.info(f"📸 Factor de ajuste visual: {factor_visual:.2f}")

            # Aplicar ajuste a coherencia lógica (las imágenes afectan la presentación)
            puntajes_criterios["coherencia_logica"]["proporcion"] *= factor_visual
            puntajes_criterios["coherencia_logica"]["proporcion"] = min(1.0, puntajes_criterios["coherencia_logica"][
                "proporcion"])

            log.info(f"  coherencia_logica ajustada: {puntajes_criterios['coherencia_logica']['proporcion']:.3f}")

        return puntajes_criterios

    except Exception as e:
        log.exception(f"Error en análisis completo: {e}")
        return {
            "aplicacion_conceptos": {"nivel": "Regular", "confidence": 0.5, "puntaje": 10.0, "proporcion": 0.5},
            "relacion_contextual": {"nivel": "Regular", "confidence": 0.5, "puntaje": 10.0, "proporcion": 0.5},
            "coherencia_logica": {"nivel": "Regular", "confidence": 0.5, "puntaje": 10.0, "proporcion": 0.5}
        }

def crear_resultado_error() -> dict:
    """
    Crea un resultado por defecto en caso de error.
    """
    return {
        "nivel": "Regular",
        "confidence": 0.5,
        "predicted_class": 2,
        "error": True
    }

# ============================================================================
# ENDPOINTS: Básicos
# ============================================================================

@app.get("/")
async def root():
    """Endpoint raíz de salud."""
    return {"status": "ok", "message": "Bienvenido a la API de EvalIA"}


@app.post("/generate-upload-url")
async def generate_upload_url(request: GenerateUploadURLRequest):
    """Genera URL firmada v4 para upload directo a GCS."""
    log.info(f"Generando URL firmada para: {request.filename}")

    try:
        blob_name = f"{uuid.uuid4()}-{request.filename}"
        bucket = storage_client.bucket(BUCKET_NAME)
        blob = bucket.blob(blob_name)

        current_time = datetime.datetime.now(timezone.utc)
        expiration = current_time + datetime.timedelta(minutes=15)

        service_account_email = os.environ.get("SERVICE_ACCOUNT_EMAIL")
        if not service_account_email:
            raise HTTPException(status_code=500, detail="SERVICE_ACCOUNT_EMAIL no configurado")

        canonical_uri = f"/{BUCKET_NAME}/{blob_name}"
        request_timestamp = current_time.strftime('%Y%m%dT%H%M%SZ')
        datestamp = current_time.strftime('%Y%m%d')
        credential_scope = f"{datestamp}/auto/storage/goog4_request"

        canonical_headers = (
            f"host:storage.googleapis.com\n"
            f"x-goog-date:{request_timestamp}\n"
        )
        signed_headers = "host;x-goog-date"

        canonical_query_string = (
            f"X-Goog-Algorithm=GOOG4-RSA-SHA256&"
            f"X-Goog-Credential={urllib.parse.quote(service_account_email + '/' + credential_scope, safe='')}&"
            f"X-Goog-Date={request_timestamp}&"
            f"X-Goog-Expires=900&"
            f"X-Goog-SignedHeaders={signed_headers}"
        )

        canonical_request = (
            f"PUT\n"
            f"{canonical_uri}\n"
            f"{canonical_query_string}\n"
            f"{canonical_headers}\n"
            f"{signed_headers}\n"
            f"UNSIGNED-PAYLOAD"
        )

        canonical_request_hash = hashlib.sha256(canonical_request.encode()).hexdigest()
        string_to_sign = (
            f"GOOG4-RSA-SHA256\n"
            f"{request_timestamp}\n"
            f"{credential_scope}\n"
            f"{canonical_request_hash}"
        )

        signing_request = iam_credentials_v1.SignBlobRequest(
            name=f"projects/-/serviceAccounts/{service_account_email}",
            payload=string_to_sign.encode('utf-8')
        )

        signing_response = iam_client.sign_blob(request=signing_request)
        signature = base64.b64encode(signing_response.signed_blob).decode('utf-8')

        signed_url = (
            f"https://storage.googleapis.com{canonical_uri}?"
            f"{canonical_query_string}&"
            f"X-Goog-Signature={urllib.parse.quote(signature, safe='')}"
        )

        log.info(f"URL firmada generada: {blob_name}")

        return {
            "url": signed_url,
            "gcs_filename": blob_name,
            "expires_at": expiration.isoformat()
        }

    except Exception as e:
        log.exception(f"Error generando URL firmada: {e}")
        raise HTTPException(status_code=500, detail=f"Error: {str(e)}")


@app.post("/upload-file-proxy")
async def upload_file_proxy(file: UploadFile = File(...), filename: str = Form(...)):
    """Proxy para subir archivos al backend."""
    log.info(f"Upload proxy para: {filename}")

    try:
        unique_filename = f"{uuid.uuid4()}-{filename}"
        file_content = await file.read()
        log.info(f"Archivo leído: {len(file_content)} bytes")

        bucket = storage_client.bucket(BUCKET_NAME)
        blob = bucket.blob(unique_filename)

        content_type = file.content_type or 'application/pdf'
        blob.upload_from_string(file_content, content_type=content_type)

        log.info(f"✅ Subido a GCS: {unique_filename}")

        return {
            "success": True,
            "filename": unique_filename,
            "gcs_uri": f"gs://{BUCKET_NAME}/{unique_filename}",
            "size": len(file_content)
        }

    except Exception as e:
        log.error(f"❌ Error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error: {str(e)}")


# ============================================================================
# ENDPOINT PRINCIPAL: Orquestación del procesamiento
# ============================================================================

@app.post("/enqueue-exam-batch")
async def enqueue_exam_batch(payload: ExamBatchRequest, db: Session = Depends(get_db)):
    """
    Orquesta el procesamiento de documentos.
    Soporta dos tipos:
    - tipo_documento='examen': Exámenes manuscritos (OCR + fuzzy matching)
    - tipo_documento='ensayo': Ensayos/Informes (extracción directa + análisis de imágenes)

    FORMATOS SOPORTADOS: PDF, DOCX
    """
    log.info(f"=== Iniciando /enqueue-exam-batch - tipo: {payload.tipo_documento} ===")

    # Validar variables de entorno
    missing_vars = []
    if not SERVICE_URL:
        missing_vars.append("SERVICE_URL")
    if not BUCKET_NAME:
        missing_vars.append("BUCKET_NAME")
    if not QUEUE_NAME:
        missing_vars.append("QUEUE_NAME")

    if missing_vars:
        log.error(f"Faltan variables: {missing_vars}")
        raise HTTPException(status_code=500, detail=f"Faltan variables: {', '.join(missing_vars)}")

    try:
        tipo_documento = payload.tipo_documento.lower()

        # ============================================================
        # VALIDACIONES ESPECÍFICAS POR TIPO
        # ============================================================

        if tipo_documento == "examen":
            # Para exámenes: REQUIERE lista de estudiantes
            if not payload.student_list or not payload.student_list.strip():
                raise HTTPException(status_code=400, detail="La lista de estudiantes está vacía")

            student_list_lines = [line.strip() for line in payload.student_list.strip().split('\n') if line.strip()]
            if not student_list_lines:
                raise HTTPException(status_code=400, detail="Lista de estudiantes no contiene nombres válidos")

            log.info(f"Lista de {len(student_list_lines)} estudiantes: {student_list_lines}")

        elif tipo_documento == "ensayo/informe":
            # Para ensayos/informes: Los nombres vienen del filename
            log.info("Tipo ensayo/informe - nombres extraídos de archivos")
            student_list_lines = None  # No se usa

        else:
            raise HTTPException(status_code=400, detail=f"Tipo de documento inválido: {tipo_documento}")

        # Validar archivos
        if not payload.pdf_files or len(payload.pdf_files) == 0:
            raise HTTPException(status_code=400, detail="No se proporcionaron archivos")

        log.info(f"Recibidos {len(payload.pdf_files)} archivos para procesar")

        # ============================================================
        # DESCARGAR Y LEER ARCHIVOS
        # ============================================================

        pdf_data_list = []
        for pdf_info in payload.pdf_files:
            blob = storage_client.bucket(BUCKET_NAME).blob(pdf_info.gcs_filename)
            if not blob.exists():
                log.error(f"Blob {pdf_info.gcs_filename} NO EXISTE")
                raise HTTPException(status_code=404, detail=f"Archivo no encontrado: {pdf_info.gcs_filename}")

            try:
                file_bytes = blob.download_as_bytes(timeout=60)
                log.info(f"Blob {pdf_info.gcs_filename} descargado: {len(file_bytes)} bytes")
            except Exception as download_err:
                log.exception(f"Error descargando {pdf_info.gcs_filename}: {download_err}")
                raise HTTPException(status_code=500, detail=f"Error descargando: {pdf_info.gcs_filename}")

            # ✅ DETECTAR TIPO DE ARCHIVO
            filename_lower = pdf_info.original_filename.lower()

            if filename_lower.endswith('.pdf'):
                # Leer PDF
                try:
                    reader = PdfReader(io.BytesIO(file_bytes))
                    num_pages = len(reader.pages)
                    log.info(f"PDF {pdf_info.original_filename} tiene {num_pages} páginas")
                    file_type = "pdf"
                except Exception as read_err:
                    log.exception(f"Error leyendo PDF {pdf_info.original_filename}: {read_err}")
                    raise HTTPException(status_code=400, detail=f"PDF corrupto: {pdf_info.original_filename}")

            elif filename_lower.endswith('.docx'):
                # Word DOCX
                log.info(f"Archivo Word detectado: {pdf_info.original_filename}")
                reader = None  # Word no usa PyPDF2
                num_pages = 1  # Word se trata como 1 "página"
                file_type = "docx"

            elif filename_lower.endswith('.doc'):
                # Word antiguo (.doc) no soportado
                raise HTTPException(
                    status_code=400,
                    detail=f"Formato .doc no soportado. Convierta a .docx o PDF: {pdf_info.original_filename}"
                )

            else:
                raise HTTPException(
                    status_code=400,
                    detail=f"Formato no soportado. Use PDF o DOCX: {pdf_info.original_filename}"
                )

            pdf_data_list.append({
                "original_filename": pdf_info.original_filename,
                "gcs_filename": pdf_info.gcs_filename,
                "bytes": file_bytes,
                "reader": reader,
                "num_pages": num_pages,
                "file_type": file_type  # ✅ AÑADIDO
            })

        # ============================================================
        # LÓGICA PARA EXÁMENES MANUSCRITOS
        # ============================================================

        if tipo_documento == "examen":
            log.info("=== PROCESAMIENTO DE EXÁMENES MANUSCRITOS ===")

            # ⚠️ NOTA: Exámenes deben ser PDF (DOCX no tiene sentido para manuscritos)
            # Validar que todos sean PDF
            for pdf_data in pdf_data_list:
                if pdf_data["file_type"] != "pdf":
                    raise HTTPException(
                        status_code=400,
                        detail=f"Exámenes manuscritos solo aceptan PDF: {pdf_data['original_filename']}"
                    )

            # Ordenar por número de cara
            def extract_page_number(filename):
                match = re.search(r'cara[_\s]*(\d+)', filename, re.IGNORECASE)
                return int(match.group(1)) if match else 999

            pdf_data_list.sort(key=lambda x: extract_page_number(x["original_filename"]))

            # ... resto del código de exámenes igual ...
            # (No necesita cambios, ya filtramos arriba)

            all_pages_face1 = []
            all_pages_face2 = []

            for idx, pdf_data in enumerate(pdf_data_list):
                face_num = idx + 1
                reader = pdf_data["reader"]
                num_pages = pdf_data["num_pages"]

                if face_num == 1:
                    all_pages_face1 = [reader.pages[i] for i in range(num_pages)]
                elif face_num == 2:
                    all_pages_face2 = [reader.pages[i] for i in range(num_pages)]

            # ... resto del código de exámenes sin cambios ...

        # ============================================================
        # LÓGICA PARA ENSAYOS/INFORMES
        # ============================================================

        elif tipo_documento == "ensayo/informe":
            log.info("=== PROCESAMIENTO DE ENSAYOS/INFORMES ===")

            evaluacion_ids = []

            # Cada archivo PDF/DOCX es un ensayo/informe de un estudiante
            for pdf_data in pdf_data_list:
                original_filename = pdf_data["original_filename"]
                gcs_filename = pdf_data["gcs_filename"]
                file_bytes = pdf_data["bytes"]
                file_type = pdf_data["file_type"]  # ✅ USAR TIPO

                log.info(f"Procesando ensayo: {original_filename} (tipo: {file_type})")

                # Extraer nombre del estudiante del filename
                student_name = original_filename
                # Remover extensión
                student_name = re.sub(r'\.(pdf|docx)$', '', student_name, flags=re.IGNORECASE)
                # Remover palabras comunes
                student_name = re.sub(r'_(ensayo|informe|trabajo|final)', '', student_name, flags=re.IGNORECASE)
                # Reemplazar guiones bajos con espacios
                student_name = student_name.replace('_', ' ').strip()

                if not student_name or len(student_name) < 3:
                    student_name = f"Estudiante_{original_filename}"

                log.info(f"Nombre extraído: {student_name}")

                # Crear evaluación
                try:
                    nueva_evaluacion = Evaluacion(
                        nombre_alumno=student_name,
                        nombre_curso=payload.nombre_curso,
                        codigo_curso=payload.codigo_curso,
                        instructor=payload.instructor,
                        semestre=payload.semestre,
                        tema=payload.tema,
                        descripcion_tema=payload.descripcion_tema,
                        tipo_documento = payload.tipo_documento
                    )
                    db.add(nueva_evaluacion)
                    db.commit()
                    db.refresh(nueva_evaluacion)
                    evaluacion_id = nueva_evaluacion.id
                    evaluacion_ids.append(evaluacion_id)
                    log.info(f"Evaluación creada ID: {evaluacion_id} para {student_name}")
                except Exception as db_err:
                    log.exception(f"Error creando evaluación: {db_err}")
                    db.rollback()
                    raise HTTPException(status_code=500, detail="Error en BD")

                # ✅ SUBIR ARCHIVO CON EXTENSIÓN CORRECTA
                # Detectar extensión original
                extension = ".pdf" if file_type == "pdf" else ".docx"
                unique_filename = f"eval_{evaluacion_id}_{original_filename}"

                # Determinar content_type correcto
                content_type = "application/pdf" if file_type == "pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

                blob_ensayo = storage_client.bucket(BUCKET_NAME).blob(unique_filename)
                blob_ensayo.upload_from_string(file_bytes, content_type=content_type, timeout=60)
                log.info(f"Ensayo subido: {unique_filename}")

                # Borrar blob original
                try:
                    blob_original = storage_client.bucket(BUCKET_NAME).blob(gcs_filename)
                    blob_original.delete(timeout=30)
                    log.info(f"Blob original borrado: {gcs_filename}")
                except Exception as del_err:
                    log.warning(f"No se pudo borrar {gcs_filename}: {del_err}")

                # Crear tarea Cloud Tasks para procesar el ensayo completo
                task_payload_dict = {
                    "gcs_filename": unique_filename,
                    "original_filename": original_filename,
                    "evaluacion_id": evaluacion_id,
                    "precomputed_ocr_text": None,
                    "tipo_documento": "ensayo/informe",
                    "file_type": file_type  # ✅ PASAR TIPO AL WORKER
                }

                task_payload_json = json.dumps(task_payload_dict)
                parent = tasks_client.queue_path(GCP_PROJECT_ID, GCP_LOCATION, QUEUE_NAME)
                task_request = tasks_v2.CreateTaskRequest(
                    parent=parent,
                    task=tasks_v2.Task(
                        http_request=tasks_v2.HttpRequest(
                            http_method=tasks_v2.HttpMethod.POST,
                            url=f"{SERVICE_URL}/process-file-task",
                            headers={"Content-Type": "application/json"},
                            body=task_payload_json.encode()
                        )
                    )
                )

                try:
                    task_response = tasks_client.create_task(request=task_request)
                    log.info(f"Tarea creada para ensayo: {task_response.name}")
                except Exception as task_err:
                    log.exception(f"Error creando tarea: {task_err}")
                    raise HTTPException(status_code=500, detail="Error creando tarea")

        # ============================================================
        # CREAR TAREAS FINALES DE EVALUACIÓN
        # ============================================================

        for evaluacion_id in evaluacion_ids:
            eval_task_payload = {"evaluacion_id": evaluacion_id}
            eval_task_json = json.dumps(eval_task_payload)

            parent = tasks_client.queue_path(GCP_PROJECT_ID, GCP_LOCATION, QUEUE_NAME)
            eval_task_request = tasks_v2.CreateTaskRequest(
                parent=parent,
                task=tasks_v2.Task(
                    http_request=tasks_v2.HttpRequest(
                        http_method=tasks_v2.HttpMethod.POST,
                        url=f"{SERVICE_URL}/process-evaluation-task",
                        headers={"Content-Type": "application/json"},
                        body=eval_task_json.encode()
                    )
                )
            )

            try:
                eval_task_response = tasks_client.create_task(request=eval_task_request)
                log.info(f"Tarea evaluación final creada: {eval_task_response.name}")
            except Exception as eval_task_err:
                log.exception(f"Error creando tarea evaluación: {eval_task_err}")

        log.info(f"=== Procesamiento completado. IDs: {evaluacion_ids} ===")
        return {
            "status": "success",
            "message": f"Se crearon {len(evaluacion_ids)} evaluaciones",
            "evaluacion_ids": evaluacion_ids,
            "tipo_documento": tipo_documento
        }

    except HTTPException as http_exc:
        log.error(f"HTTP Exception: {http_exc.status_code} - {http_exc.detail}")
        raise http_exc
    except Exception as e:
        log.exception(f"ERROR CRÍTICO: {e}")
        raise HTTPException(status_code=500, detail="Error interno inesperado")


# ============================================================================
# ENDPOINT WORKER: Procesa un archivo (con lógica diferenciada)
# ============================================================================

@app.post("/process-file-task")
async def process_file_task(payload: FileTaskPayload, db: Session = Depends(get_db)):
    """
    WORKER: Procesa un solo archivo.
    - Para exámenes: usa OCR
    - Para ensayos: extrae texto directo + analiza imágenes con Gemini

    FORMATOS SOPORTADOS: PDF, DOCX
    """
    log.info(f"[TASK-FILE] Iniciando para {payload.original_filename} - tipo: {payload.tipo_documento}")

    texto_extraido = payload.precomputed_ocr_text
    analisis_visual = None
    blob = storage_client.bucket(BUCKET_NAME).blob(payload.gcs_filename)
    archivo_id = None

    # ✅ DETECTAR TIPO DE ARCHIVO
    file_type = getattr(payload, 'file_type', None)
    if not file_type:
        # Fallback: detectar por extensión
        filename_lower = payload.original_filename.lower()
        if filename_lower.endswith('.docx'):
            file_type = "docx"
        elif filename_lower.endswith('.pdf'):
            file_type = "pdf"
        else:
            file_type = "pdf"  # Default

    log.info(f"[TASK-FILE] Tipo de archivo: {file_type}")

    try:
        if not blob.exists():
            log.error(f"[TASK-FILE] Blob {payload.gcs_filename} NO EXISTE")
            raise HTTPException(status_code=404, detail=f"Archivo no encontrado: {payload.gcs_filename}")

        if not texto_extraido:
            log.info(f"[TASK-FILE] Descargando y extrayendo texto...")
            try:
                file_bytes = blob.download_as_bytes(timeout=60)
            except Exception as download_err:
                log.exception(f"[TASK-FILE] Error descargando: {download_err}")
                raise HTTPException(status_code=500, detail="Error descargando archivo")

            log.info(f"[TASK-FILE] Blob descargado: {len(file_bytes)} bytes")

            # ============================================================
            # LÓGICA DIFERENCIADA POR TIPO
            # ============================================================

            if payload.tipo_documento == "examen":
                # EXAMEN MANUSCRITO: Usar OCR (solo PDF)
                log.info(f"[TASK-FILE] Aplicando OCR para examen manuscrito...")
                texto_extraido = call_ocr_api(file_bytes, "application/pdf")
                log.info(f"[TASK-FILE] OCR completado: {len(texto_extraido) if texto_extraido else 0} chars")

            elif payload.tipo_documento == "ensayo/informe":
                # ENSAYO/INFORME: Extracción según formato
                log.info(f"[TASK-FILE] Extracción para ensayo/informe (formato: {file_type})...")

                try:
                    if file_type == "docx":
                        # ✅ EXTRAER TEXTO DE WORD
                        texto_extraido = extraer_texto_word(file_bytes)
                        log.info(f"[TASK-FILE] Texto extraído de Word: {len(texto_extraido)} chars")

                        # ✅ EXTRAER IMÁGENES DE WORD
                        imagenes_docx = extract_images_from_docx(file_bytes)
                        deteccion = {
                            "texto": texto_extraido,
                            "tiene_imagenes": len(imagenes_docx) > 0,
                            "imagenes": imagenes_docx,
                            "requiere_ocr": len(texto_extraido) < 100
                        }

                    elif file_type == "pdf":
                        # ✅ EXTRAER TEXTO E IMÁGENES DE PDF
                        deteccion = detectar_tipo_y_extraer_texto(file_bytes, payload.original_filename)

                    else:
                        # Fallback
                        log.warning(f"[TASK-FILE] Tipo desconocido, usando OCR")
                        texto_extraido = call_ocr_api(file_bytes, "application/pdf")
                        deteccion = {"texto": texto_extraido, "tiene_imagenes": False, "imagenes": [],
                                     "requiere_ocr": False}

                    if deteccion.get("requiere_ocr"):
                        # Fallback si no tiene texto
                        log.warning(f"[TASK-FILE] Sin texto, aplicando OCR como fallback")
                        texto_extraido = call_ocr_api(file_bytes, "application/pdf")
                    else:
                        log.info(f"[TASK-FILE] Texto extraído exitosamente")
                        texto_extraido = deteccion["texto"]

                    # ✅ ANALIZAR IMÁGENES CON GEMINI (PDF o Word)
                    if deteccion.get("tiene_imagenes") and len(deteccion.get("imagenes", [])) > 0:
                        log.info(f"[TASK-FILE] Encontradas {len(deteccion['imagenes'])} imágenes")

                        evaluacion = db.query(Evaluacion).filter(Evaluacion.id == payload.evaluacion_id).first()
                        if evaluacion and gemini_client:
                            try:
                                analisis_visual = analyze_images_with_gemini(
                                    deteccion["imagenes"],
                                    evaluacion.tema or "No especificado",
                                    evaluacion.descripcion_tema or ""
                                )
                                log.info(f"[TASK-FILE] Análisis visual completado con Gemini")
                            except Exception as vision_err:
                                log.exception(f"[TASK-FILE] Error análisis visual: {vision_err}")
                                analisis_visual = {"error": "No se pudo analizar"}
                        else:
                            log.warning(f"[TASK-FILE] Gemini no disponible, saltando análisis visual")

                except Exception as deteccion_err:
                    log.exception(f"[TASK-FILE] Error detección, fallback OCR: {deteccion_err}")
                    texto_extraido = call_ocr_api(file_bytes, "application/pdf")

            else:
                log.warning(f"[TASK-FILE] Tipo desconocido: {payload.tipo_documento}, usando OCR")
                texto_extraido = call_ocr_api(file_bytes, "application/pdf")

        else:
            log.info(f"[TASK-FILE] Usando texto precalculado")

        # ============================================================
        # GUARDAR EN BASE DE DATOS
        # ============================================================

        try:
            archivo_existente = db.query(ArchivoProcesado) \
                .filter(ArchivoProcesado.evaluacion_id == payload.evaluacion_id,
                        ArchivoProcesado.nombre_archivo_original == payload.original_filename) \
                .with_for_update().first()

            if archivo_existente:
                log.warning(f"[TASK-FILE] Ya existe ArchivoProcesado ID: {archivo_existente.id}. Actualizando.")
                archivo_existente.texto_extraido = texto_extraido or ""

                # ✅ GUARDAR ANÁLISIS VISUAL COMO JSON
                if analisis_visual:
                    archivo_existente.analisis_visual = json.dumps(analisis_visual, ensure_ascii=False)
                    log.info(f"[TASK-FILE] Análisis visual guardado en BD (actualización)")

                db.commit()
                db.refresh(archivo_existente)
                archivo_id = archivo_existente.id
            else:
                # ✅ PREPARAR ANÁLISIS VISUAL COMO JSON
                analisis_visual_json = None
                if analisis_visual:
                    analisis_visual_json = json.dumps(analisis_visual, ensure_ascii=False)
                    log.info(f"[TASK-FILE] Análisis visual será guardado en BD")

                archivo_procesado = ArchivoProcesado(
                    nombre_archivo_original=payload.original_filename,
                    texto_extraido=texto_extraido or "",
                    evaluacion_id=payload.evaluacion_id,
                    analisis_visual=analisis_visual_json  # ✅ GUARDAR AQUÍ
                )
                log.info(f"[TASK-FILE] Guardando NUEVO ArchivoProcesado")
                db.add(archivo_procesado)
                db.commit()
                db.refresh(archivo_procesado)
                archivo_id = archivo_procesado.id

            log.info(f"[TASK-FILE] ArchivoProcesado guardado con ID: {archivo_id}")

        except sqlalchemy.exc.OperationalError as db_op_err:
            log.exception(f"[TASK-FILE] Error operacional BD: {db_op_err}")
            db.rollback()
            raise HTTPException(status_code=503, detail="Error temporal BD, reintentar")
        except Exception as db_err:
            log.exception(f"[TASK-FILE] Error guardando en BD: {db_err}")
            db.rollback()
            raise HTTPException(status_code=500, detail="Error guardando en BD")

    except HTTPException as http_exc:
        log.error(f"[TASK-FILE] HTTP Exception: {http_exc.status_code} - {http_exc.detail}")
        raise http_exc
    except Exception as e:
        log.exception(f"[TASK-FILE] Error INESPERADO: {e}")
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Error inesperado: {e}")
    finally:
        # ============================================================
        # LIMPIEZA: Borrar blob de GCS
        # ============================================================
        try:
            blob_ref_final = storage_client.bucket(BUCKET_NAME).blob(payload.gcs_filename)
            if blob_ref_final.exists():
                blob_ref_final.delete(timeout=30)
                log.info(f"[TASK-FILE] Blob eliminado: {payload.gcs_filename}")
            else:
                log.info(f"[TASK-FILE] Blob ya borrado")
        except Exception as e:
            log.warning(f"[TASK-FILE] No se pudo eliminar blob: {e}")

    if archivo_id:
        return {"status": "success", "archivo_id": archivo_id}
    else:
        log.error("[TASK-FILE] Finalizó sin error pero no hay archivo_id")
        raise HTTPException(status_code=500, detail="Error: No se obtuvo archivo_id")


# ============================================================================
# ENDPOINT WORKER: Procesa evaluación completa
# ============================================================================

@app.post("/process-evaluation-task")
async def process_evaluation_task(request: Request, db: Session = Depends(get_db)):
    """Procesa la evaluación final después de analizar todos los archivos."""
    try:
        data = await request.json()
        evaluacion_id = data.get("evaluacion_id")

        log.info(f"[TASK-EVAL] Iniciando para evaluacion_id: {evaluacion_id}")

        # Obtener evaluación
        evaluacion = db.query(Evaluacion).filter(Evaluacion.id == evaluacion_id).first()
        if not evaluacion:
            log.error(f"Evaluación {evaluacion_id} no encontrada")
            return {"error": "Evaluacion no encontrada"}

        # ✅ OBTENER PESOS DESDE BD
        criterio_config = db.query(CriterioConfig).first()
        if not criterio_config:
            log.warning("No hay CriterioConfig, usando valores por defecto")
            peso_aplicacion = 0.4
            peso_relacion = 0.3
            peso_coherencia = 0.3
        else:
            peso_aplicacion = criterio_config.aplicacion_conceptos
            peso_relacion = criterio_config.relacion_contextual
            peso_coherencia = criterio_config.coherencia_logica

        log.info(f"[TASK-EVAL] Pesos: Apl={peso_aplicacion}, Rel={peso_relacion}, Coh={peso_coherencia}")

        # Obtener archivos procesados
        archivos = db.query(ArchivoProcesado).filter(
            ArchivoProcesado.evaluacion_id == evaluacion_id
        ).all()

        if not archivos:
            return {"error": "No hay archivos"}

        # Concatenar texto
        texto_completo = "\n\n".join([a.texto_extraido for a in archivos if a.texto_extraido])
        log.info(f"[TASK-EVAL] Texto: {len(texto_completo)} chars")

        # ✅ RECUPERAR ANÁLISIS VISUAL DESDE LA BASE DE DATOS
        analisis_visual_acumulado = {}
        num_analisis = 0

        for archivo in archivos:
            if archivo.analisis_visual:
                try:
                    analisis = json.loads(archivo.analisis_visual)
                    num_analisis += 1

                    # Combinar análisis (promediando puntajes si hay múltiples archivos)
                    if not analisis_visual_acumulado:
                        analisis_visual_acumulado = analisis
                    else:
                        # Promediar puntajes de claridad, relevancia y calidad técnica
                        for key in ["claridad_visual", "relevancia_tema", "calidad_tecnica"]:
                            if key in analisis and key in analisis_visual_acumulado:
                                if isinstance(analisis[key], dict) and "puntaje" in analisis[key]:
                                    puntaje_actual = analisis_visual_acumulado[key].get("puntaje", 0)
                                    puntaje_nuevo = analisis[key].get("puntaje", 0)
                                    analisis_visual_acumulado[key]["puntaje"] = (puntaje_actual + puntaje_nuevo) / 2

                    log.info(f"[TASK-EVAL] Análisis visual recuperado de: {archivo.nombre_archivo_original}")
                except json.JSONDecodeError as e:
                    log.warning(f"Error parseando análisis visual de {archivo.nombre_archivo_original}: {e}")

        analisis_visual = analisis_visual_acumulado if analisis_visual_acumulado else None

        if analisis_visual:
            log.info(f"[TASK-EVAL] ✅ Análisis visual disponible ({num_analisis} archivo(s) con imágenes)")
        else:
            log.info(f"[TASK-EVAL] ⚠️ No hay análisis visual disponible")

        # ✅ ANÁLISIS ML POR CRITERIO
        resultados_criterios = analizar_texto_con_modelo(texto_completo, analisis_visual)

        # Extraer proporciones (0-1)
        aplicacion_decimal = resultados_criterios["aplicacion_conceptos"]["proporcion"]
        relacion_decimal = resultados_criterios["relacion_contextual"]["proporcion"]
        coherencia_decimal = resultados_criterios["coherencia_logica"]["proporcion"]

        log.info(
            f"[TASK-EVAL] Proporciones: Apl={aplicacion_decimal:.3f}, Rel={relacion_decimal:.3f}, Coh={coherencia_decimal:.3f}")

        # ✅ CALCULAR NOTA FINAL PONDERADA (0-20)
        # Convertir proporciones a puntajes
        puntaje_aplicacion = aplicacion_decimal * 20
        puntaje_relacion = relacion_decimal * 20
        puntaje_coherencia = coherencia_decimal * 20

        nota_final = (
                puntaje_aplicacion * peso_aplicacion +
                puntaje_relacion * peso_relacion +
                puntaje_coherencia * peso_coherencia
        )

        nota_final = max(0, min(20, nota_final))

        log.info(f"[TASK-EVAL] Nota final: {nota_final:.2f}")

        # ✅ GUARDAR EN BD
        resultado_existente = db.query(ResultadoAnalisis).filter(
            ResultadoAnalisis.evaluacion_id == evaluacion_id
        ).first()

        if resultado_existente:
            resultado_existente.aplicacion_conceptos = float(aplicacion_decimal)
            resultado_existente.relacion_contextual = float(relacion_decimal)
            resultado_existente.coherencia_logica = float(coherencia_decimal)
            resultado_existente.nota_final = float(nota_final)
            db.commit()
            log.info(f"[TASK-EVAL] Resultado actualizado")
        else:
            nuevo_resultado = ResultadoAnalisis(
                evaluacion_id=evaluacion_id,
                aplicacion_conceptos=float(aplicacion_decimal),
                relacion_contextual=float(relacion_decimal),
                coherencia_logica=float(coherencia_decimal),
                nota_final=float(nota_final)
            )
            db.add(nuevo_resultado)
            db.commit()
            log.info(f"[TASK-EVAL] Resultado creado")

        # Actualizar estado
        evaluacion.estado = "completada"
        db.commit()

        return {
            "status": "completed",
            "evaluacion_id": evaluacion_id,
            "nota_final": nota_final,
            "aplicacion_conceptos": aplicacion_decimal,
            "relacion_contextual": relacion_decimal,
            "coherencia_logica": coherencia_decimal,
            "detalles": resultados_criterios,
            "analisis_visual_usado": analisis_visual is not None
        }

    except Exception as e:
        log.exception(f"Error en process_evaluation_task: {e}")
        db.rollback()
        return {"error": str(e)}

# ============================================================================
# ENDPOINTS: Consulta de evaluaciones
# ============================================================================

@app.get("/evaluaciones")
async def get_evaluaciones(db: Session = Depends(get_db)):
    """Obtiene lista de evaluaciones."""
    try:
        evaluaciones = db.query(Evaluacion) \
            .options(joinedload(Evaluacion.resultado_analisis)) \
            .order_by(Evaluacion.fecha_creacion.desc()) \
            .limit(1000) \
            .all()

        return [EvaluacionSchema.from_orm(ev) for ev in evaluaciones]
    except Exception as e:
        log.exception(f"Error obteniendo evaluaciones: {e}")
        raise HTTPException(status_code=500, detail="Error obteniendo evaluaciones")


@app.get("/evaluaciones/{evaluacion_id}")
async def get_evaluacion_detail(evaluacion_id: int, db: Session = Depends(get_db)):
    """Obtiene detalle completo de una evaluación."""
    try:
        evaluacion = db.query(Evaluacion) \
            .options(joinedload(Evaluacion.resultado_analisis), joinedload(Evaluacion.archivos_procesados)) \
            .filter(Evaluacion.id == evaluacion_id) \
            .first()

        if not evaluacion:
            raise HTTPException(status_code=404, detail="Evaluación no encontrada")

        return EvaluacionDetailSchema.from_orm(evaluacion)
    except HTTPException:
        raise
    except Exception as e:
        log.exception(f"Error obteniendo evaluación {evaluacion_id}: {e}")
        raise HTTPException(status_code=500, detail="Error obteniendo evaluación")


# ============================================================================
# ENDPOINTS: Configuración de criterios
# ============================================================================

@app.get("/criterios")
async def get_criterios(db: Session = Depends(get_db)):
    """Obtiene configuración de criterios."""
    log.info("GET /criterios")
    try:
        config = db.query(CriterioConfig).first()
        if not config:
            config = CriterioConfig(
                aplicacion_conceptos=0.4,
                relacion_contextual=0.3,
                coherencia_logica=0.3
            )
            db.add(config)
            db.commit()
            db.refresh(config)

        log.info(f"Criterios obtenidos: {config.__dict__}")
        return {
            "aplicacion_conceptos": config.aplicacion_conceptos,
            "relacion_contextual": config.relacion_contextual,
            "coherencia_logica": config.coherencia_logica
        }
    except Exception as e:
        log.exception(f"Error obteniendo criterios: {e}")
        raise HTTPException(status_code=500, detail="Error obteniendo criterios")


@app.post("/criterios")
async def update_criterios(payload: CriterioConfigUpdate, db: Session = Depends(get_db)):
    """Actualiza configuración de criterios."""
    log.info(f"POST /criterios: {payload.dict()}")
    try:
        suma = payload.aplicacion_conceptos + payload.relacion_contextual + payload.coherencia_logica
        if abs(suma - 1.0) > 0.01:
            raise HTTPException(status_code=400, detail="La suma debe ser 1.0")

        config = db.query(CriterioConfig).first()
        if not config:
            config = CriterioConfig()
            db.add(config)

        config.aplicacion_conceptos = payload.aplicacion_conceptos
        config.relacion_contextual = payload.relacion_contextual
        config.coherencia_logica = payload.coherencia_logica

        db.commit()
        db.refresh(config)

        log.info(f"Criterios actualizados")
        return {
            "status": "success",
            "criterios": {
                "aplicacion_conceptos": config.aplicacion_conceptos,
                "relacion_contextual": config.relacion_contextual,
                "coherencia_logica": config.coherencia_logica
            }
        }
    except HTTPException:
        raise
    except Exception as e:
        log.exception(f"Error actualizando criterios: {e}")
        db.rollback()
        raise HTTPException(status_code=500, detail="Error actualizando criterios")


# ============================================================================
# ENDPOINTS: Transcripciones (SOLO PARA EXÁMENES)
# ============================================================================

@app.get("/evaluaciones/{evaluacion_id}/transcripcion/pdf")
async def get_transcripcion_pdf(evaluacion_id: int, db: Session = Depends(get_db)):
    """
    Genera PDF con transcripción OCR.
    SOLO disponible para exámenes manuscritos.
    """
    try:
        evaluacion = db.query(Evaluacion).filter(Evaluacion.id == evaluacion_id).first()
        if not evaluacion:
            raise HTTPException(status_code=404, detail="Evaluación no encontrada")

        archivos = db.query(ArchivoProcesado) \
            .filter(ArchivoProcesado.evaluacion_id == evaluacion_id) \
            .all()

        if not archivos:
            raise HTTPException(status_code=404, detail="No hay archivos procesados")

        # Generar PDF
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        width, height = letter

        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, height - 50, f"Transcripción - {evaluacion.nombre_alumno}")

        c.setFont("Helvetica", 10)
        y_position = height - 100

        for idx, archivo in enumerate(archivos, 1):
            if y_position < 100:
                c.showPage()
                y_position = height - 50

            c.setFont("Helvetica-Bold", 12)
            c.drawString(50, y_position, f"Página {idx}:")
            y_position -= 20

            c.setFont("Helvetica", 9)
            texto = archivo.texto_extraido or "[Sin texto]"

            for line in texto.split('\n'):
                if y_position < 50:
                    c.showPage()
                    y_position = height - 50
                c.drawString(60, y_position, line[:100])
                y_position -= 15

            y_position -= 20

        c.save()
        buffer.seek(0)

        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename=transcripcion_{evaluacion.nombre_alumno.replace(' ', '_')}.pdf"
            }
        )

    except HTTPException:
        raise
    except Exception as e:
        log.exception(f"Error generando PDF: {e}")
        raise HTTPException(status_code=500, detail="Error generando PDF")


@app.post("/evaluaciones/transcripciones/zip")
async def get_transcripciones_zip(evaluacion_ids: List[int] = Body(...), db: Session = Depends(get_db)):
    """
    Descarga múltiples transcripciones en ZIP.
    SOLO para exámenes manuscritos.
    """
    try:
        evaluaciones = db.query(Evaluacion).filter(Evaluacion.id.in_(evaluacion_ids)).all()
        if not evaluaciones:
            raise HTTPException(status_code=404, detail="No se encontraron evaluaciones")

        zip_buffer = io.BytesIO()

        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
            for evaluacion in evaluaciones:
                archivos = db.query(ArchivoProcesado) \
                    .filter(ArchivoProcesado.evaluacion_id == evaluacion.id) \
                    .all()

                if not archivos:
                    continue

                # Generar PDF para esta evaluación
                pdf_buffer = io.BytesIO()
                c = canvas.Canvas(pdf_buffer, pagesize=letter)
                width, height = letter

                c.setFont("Helvetica-Bold", 16)
                c.drawString(50, height - 50, f"Transcripción - {evaluacion.nombre_alumno}")

                y_position = height - 100

                for idx, archivo in enumerate(archivos, 1):
                    if y_position < 100:
                        c.showPage()
                        y_position = height - 50

                    c.setFont("Helvetica-Bold", 12)
                    c.drawString(50, y_position, f"Página {idx}:")
                    y_position -= 20

                    c.setFont("Helvetica", 9)
                    texto = archivo.texto_extraido or "[Sin texto]"

                    for line in texto.split('\n'):
                        if y_position < 50:
                            c.showPage()
                            y_position = height - 50
                        c.drawString(60, y_position, line[:100])
                        y_position -= 15

                    y_position -= 20

                c.save()
                pdf_buffer.seek(0)

                filename = f"transcripcion_{evaluacion.nombre_alumno.replace(' ', '_')}_{evaluacion.id}.pdf"
                zip_file.writestr(filename, pdf_buffer.getvalue())

        zip_buffer.seek(0)

        return StreamingResponse(
            zip_buffer,
            media_type="application/zip",
            headers={"Content-Disposition": "attachment; filename=transcripciones.zip"}
        )

    except HTTPException:
        raise
    except Exception as e:
        log.exception(f"Error generando ZIP: {e}")
        raise HTTPException(status_code=500, detail="Error generando ZIP")
